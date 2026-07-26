from pathlib import Path
from pypdf import PdfReader
from docx import Document
import easyocr
import warnings

warnings.filterwarnings("ignore")

reader = easyocr.Reader(
    ['en'],
    gpu=False,
    verbose=False
)


def _read_pdf_with_pymupdf(filepath):
    """Most tolerant option — handles malformed/oddly-encoded PDFs well."""
    import fitz  # PyMuPDF

    text = ""
    doc = fitz.open(filepath)

    for page in doc:
        page_text = page.get_text()
        if page_text:
            text += page_text + "\n"

    doc.close()
    return text


def _read_pdf_with_pypdf(filepath):
    """Original approach, run in lenient (non-strict) mode."""
    pdf = PdfReader(filepath, strict=False)
    text = ""

    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text


def _read_pdf_with_ocr(filepath):
    """Last resort — rasterize each page and OCR it (handles scanned
    or otherwise unparsable PDFs)."""
    import fitz

    text = ""
    doc = fitz.open(filepath)

    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        temp_img = Path(filepath).with_suffix(f".page{i}.png")
        pix.save(temp_img)

        try:
            page_text = read_image(temp_img)
            if page_text:
                text += page_text + "\n"
        finally:
            temp_img.unlink(missing_ok=True)

    doc.close()
    return text


def read_pdf(filepath):

    errors = []

    # 1. Try PyMuPDF first — most robust against malformed PDFs
    try:
        text = _read_pdf_with_pymupdf(filepath)
        if text.strip():
            return text
    except Exception as e:
        errors.append(f"PyMuPDF: {e}")

    # 2. Fall back to pypdf in lenient mode
    try:
        text = _read_pdf_with_pypdf(filepath)
        if text.strip():
            return text
    except Exception as e:
        errors.append(f"pypdf: {e}")

    # 3. Last resort — OCR each page as an image (handles scanned PDFs
    #    or files with no extractable text layer)
    try:
        text = _read_pdf_with_ocr(filepath)
        if text.strip():
            return text
    except Exception as e:
        errors.append(f"OCR: {e}")

    raise Exception(
        "Unable to read PDF. The file may be corrupted or password "
        "protected. Details: " + " | ".join(errors)
    )


def read_docx(filepath):
    document = Document(filepath)
    text = ""

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_image(filepath):
    result = reader.readtext(str(filepath), detail=0)
    return "\n".join(result)


def extract_text(filepath):
    filepath = Path(filepath)

    extension = filepath.suffix.lower()

    if extension == ".pdf":
        return read_pdf(filepath)

    elif extension == ".docx":
        return read_docx(filepath)

    elif extension in [
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".tiff"
    ]:
        return read_image(filepath)

    else:
        raise ValueError("Unsupported file format.")
